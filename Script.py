from docx import Document
import os

def iniciar_documento_nuevo(directorio_destino, contador_archivos, texto_actual):
    """
    Inicia un nuevo documento con el texto acumulado y actualiza el contador de archivos.
    """
    ruta_nuevo_archivo = os.path.join(directorio_destino, f'separado_{contador_archivos}.docx')
    doc_nuevo = Document()
    doc_nuevo.add_paragraph(texto_actual)
    doc_nuevo.save(ruta_nuevo_archivo)
    return contador_archivos + 1

def es_titulo(parrafo, nombre_fuente='Arial', tamano_fuente_min=21, tamano_fuente_max=25):
    """
    Evalúa si el parrafo completo se considera un título, basado en el formato del primer run.
    """
    if not parrafo.runs:  # Si no hay runs, no es un título
        return False
    run = parrafo.runs[0]
    fuente = run.font.name
    tamano_fuente = run.font.size
    tamano_en_pt = tamano_fuente.pt if tamano_fuente else None

    # Añadir una comprobación para asegurar que tamano_en_pt no sea None
    if tamano_en_pt is None:
        return False

    return fuente == nombre_fuente and tamano_fuente_min <= tamano_en_pt <= tamano_fuente_max
  
def separar_documento(ruta_documento, directorio_destino):
    """
    Separa el documento en varios documentos basados en títulos específicos, considerando
    títulos que ocupan múltiples líneas.
    """
    os.makedirs(directorio_destino, exist_ok=True)
    doc = Document(ruta_documento)
    texto_actual = ""
    contador_archivos = 1
    en_titulo = False

    for parrafo in doc.paragraphs:
        if es_titulo(parrafo):
            if en_titulo:  # Continúa el título anterior
                texto_actual += parrafo.text + "\n"
            else:  # Nuevo título
                if texto_actual:  # Guarda la sección anterior si existe
                    contador_archivos = iniciar_documento_nuevo(directorio_destino, contador_archivos, texto_actual)
                    texto_actual = ""
                texto_actual += parrafo.text + "\n"
                en_titulo = True
        else:
            if en_titulo:  # El primer parágrafo después del título, marca el inicio del contenido
                en_titulo = False
            texto_actual += parrafo.text + "\n"

    # Guardar la última sección después del último título
    if texto_actual:
        iniciar_documento_nuevo(directorio_destino, contador_archivos, texto_actual)

# Uso del código
ruta_documento = 'direction.docx'
directorio_destino = 'direction_folder'
separar_documento(ruta_documento, directorio_destino)
